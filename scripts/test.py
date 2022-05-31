from data_utils import ASSETS_PATH
from docx import Document
import time
import glob
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

def get_table_data(output_doc_name, table):
    """
    获取表格数据
    """
    paragraph = output_doc_name.add_paragraph()
    paragraph._p.addnext(table._tbl)

def is_cut_doc(paragraph,head,tail,is_cut=False,is_final=False):
    """
    判断是否是切分文档
    """
    if head in paragraph.text:
        print('有开头')
        is_cut = True
        return is_cut,paragraph,is_final
    if tail in paragraph.text:
        print('有结尾')
        is_cut = False
        is_final = True
        paragraph = Paragraph()
        paragraph.text = ''
        return is_cut,paragraph,is_final
    if is_cut == True:
        print('有中间')
        return is_cut,paragraph,is_final

def get_para_data(output_doc_name, paragraph):
    """
    获取段落数据
    """
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
        # Run's font size
        output_run.font.size = run.font.size
    # Paragraph's alignment data
    output_para.style = paragraph.style
    output_para.alignment = paragraph.alignment
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control


cuts = [
    {'name':'投标函及投标函附录'
    ,'start':'投标函及投标函附录'
    ,'end':'法定代表人身份证明'},
]

files = glob.glob(ASSETS_PATH + '*.docx')

for file in files:
    document = Document(file)
    for cut in cuts:
        head = cut['start']
        tail = cut['end']
        start = False
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, document)
                if head in paragraph.text:
                    new_doc = Document()
                    get_para_data(new_doc,paragraph)
                    start = True
                if tail in paragraph.text:
                    new_doc.save(cut['name'] + str(int(time.time()*100000))  + '.docx') 
                    start = False
                if start == True:
                    get_para_data(new_doc,paragraph) 
            elif isinstance(child, CT_Tbl):
                if start == True:
                    table = Table(child, document)
                    get_table_data(new_doc,table)